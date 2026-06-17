import docx
import pandas as pd
import re

"""
This function will take the path and name of a docx and write it to a txt.
"""


class ProcessCyGr:
    """
    Inputs:
      - path: should be that path to find the rules
      - rules_files: a list of excel documents that contains an ordered
                      list of the files that contain the rules.
    """

    def __init__(self, data_path: str, rule_path: str, rule_files: list) -> None:
        self.DATA_PATH = data_path
        self.RULE_PATH = rule_path
        # converts it to a list of Pandas
        self.rules = list(
            map(lambda x: pd.read_excel(rule_path + x).fillna(""), rule_files)
        )

        # variables below are for removing the second stress
        self.exceptions = [
            "ίνταμπὄν’",
            "ίνταμπὄνι",
            "ίνταμπὀννά",
            "ίνταμπὄνουν",
            "ίνταμπὄνουσιν",
            "ίνναμπὄν’",
            "ίνναμπὄνι",
            "ίνναμπὀννά",
            "ίνναμπὄνουν",
            "ίνναμπὄνουσιν",
            "ίνταμπὄχω",
            "ίνταμπὄσ̆ει",
            "ίνταμπὄσ̆εις",
            "ίνταμπὄχουμεν",
            "ίνταμπὄσ̆ετε",
            "ίνταμπὄχουν",
            "ίνταμπὄχουσιν",
        ]
        self.accent_dict = {
            "έ": "ε",
            "ό": "ο",
            "ή": "η",
            "ώ": "ω",
            "ά": "α",
            "ί": "ι",
            "ύ": "υ",
            "ὄ": "ὀ",
            "ΐ": "ϊ",
            "ΰ": "ϋ",
            "Έ": "Ε",
            "Ό": "Ο",
            "Ή": "Η",
            "Ώ": "Ω",
            "Ά": "Α",
            "Ί": "Ι",
            "Ύ": "Υ",
            "Ὄ": "Ὀ",
        }

    def __exceptions_match__(self, word) -> bool:
        for exception in self.exceptions:
            if exception in word:
                return True
        return False

    def __check_length__(self, infile, outfile) -> bool:
        return len(infile) == len(outfile)

    def __read_word_file__(self, filename):
        doc = docx.Document(self.DATA_PATH + filename)
        lines = []
        for para in doc.paragraphs:
            lines.append(para.text)
        return lines

    def __read_txt_file__(self, filename):
        f = open(self.DATA_PATH + filename, "r")
        lines = []
        for l in f:
            lines.append(l)
        return lines

    def __doc_to_txt__(self, filename, lines):
        fullTextPar = "\n".join(lines)

        f = open(self.DATA_PATH + filename + ".txt", "w", encoding="utf-8")
        f.write(fullTextPar)
        f.close()

    def __lines_to_doc__(self, filename, lines):
        doc = docx.Document()
        for line in lines:
            p = doc.add_paragraph()
            p.add_run(line)
            p = doc.add_paragraph()
        doc.save(self.DATA_PATH + filename + ".docx")

    def __stress_helper__(self, word):
        fixed_word = []
        split_char = "-"
        if "–" in word:
            split_char = "–"
        for curr_word in word.split(split_char):
            accents = list(
                filter(lambda letter: letter in self.accent_dict.keys(), curr_word)
            )
            # print(word, not self.__exceptions_match__(word))
            # if there's more than 1 accent (AKA two accents), we may need to remove it
            if len(accents) > 1 and not self.__exceptions_match__(curr_word):
                # if len(accents) > 1 and word not in self.exceptions:
                # use the accents list to identify the second character, and map it to its unaccented version
                # i think this has to be done by reversing the word string first (in the case that the accented
                # words found are the same character)
                # NOTE: the implementation below will definitely be expensive, look into other ways to do this
                reversed_word = curr_word[::-1]
                accent_index = reversed_word.find(accents[1])
                # WRONG -> strings are immutable -> reversed_word[accent_index] = self.accent_dict[accents[1]] # replace the key w the value
                reversed_word = (
                    reversed_word[:accent_index]
                    + self.accent_dict[accents[1]]
                    + reversed_word[accent_index + 1 :]
                )
                fixed_word.append(reversed_word[::-1])
                # return reversed_word[::-1]
            # return word
            else:
                fixed_word.append(curr_word)
        # print(fixed_word)
        return "-".join(fixed_word)

    def __remove_second_stress__(self, line):
        # Step 1: break the line into list of words
        words = line.split()

        # Step 2: fix the words for the second stress correction
        words = list(map(self.__stress_helper__, words))
        return " ".join(words)

    """
  The rules have to be applied in sequential order for each line.
  NOTE: Including the second stress removal, which has to happen
        after the smoothing and before the corrections rules. I will
        implement it as a helper function which is called here.
  """

    def __FaR_helper__(self, line):
        # apply this for each rule file
        # print(f"Line before preprocessing: {line}")
        for j, rule_list in enumerate(self.rules):
            # NOTE: this assumes we are the beginning of everything.
            if j == 0:
                # we need to add specific things at the begging and end of each line
                line = "℗♯" + line + "♯℗"
                # print(f"Adding special character before all rules: {line}")
            # NOTE: this assumes that we are right before the 'corrections' and after
            # the 'smoothing' rules.
            elif j == 1:
                # we need to remove the second stress right after smoothing document
                line = self.__remove_second_stress__(line)
                # print(f"Removing second stresses: {line}")
            find = rule_list["Find"]
            replace = rule_list["Replace"]

            for i in range(len(find)):
                line = line.replace(find[i], replace[i])

            # print(f"Current state at the end of running {j}th loop: {line}")
        return line

    def find_and_replace(self, filename, is_word_doc=True):
        # Step 1: Convert word to text doc is required
        lines = None
        if is_word_doc:
            lines = self.__read_word_file__(filename)
            self.__doc_to_txt__(filename, lines)
        else:
            lines = self.__read_txt_file__(filename)

        # Step 2: Apply the rules in appropriate order for all lines
        processed_lines = list(map(self.__FaR_helper__, lines))
        # print(processed_lines[0])
        # print(processed_lines)

        # Step 3: Save the updated lines into a txt and word doc version
        filename = filename.split(".")[0]  # parse the filename
        self.__doc_to_txt__(filename + "-normalized", processed_lines)
        self.__lines_to_doc__(filename + "-normalized", processed_lines)

    def normalize_text(self, text):
        processed_text = self.__FaR_helper__(text)
        return processed_text
